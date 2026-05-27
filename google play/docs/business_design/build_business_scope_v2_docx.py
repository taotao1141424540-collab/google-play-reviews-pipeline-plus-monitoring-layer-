#!/usr/bin/env python3
"""Write business_scope_and_metrics_v2.docx from business_scope_and_metrics_v2.md.

Requires: pip install python-docx

Run from repo root:
  python3 "google play/docs/business_design/build_business_scope_v2_docx.py"

Or from this directory:
  python3 build_business_scope_v2_docx.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

HERE = Path(__file__).resolve().parent
MD_PATH = HERE / "business_scope_and_metrics_v2.md"
DOCX_PATH = HERE / "business_scope_and_metrics_v2.docx"


def _set_run_font_cn(run, name: str = "Microsoft YaHei", size_pt: float = 10.5) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)


def add_paragraph_with_bold(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = p.add_run(part[2:-2])
            run.bold = True
            _set_run_font_cn(run)
        else:
            run = p.add_run(part)
            _set_run_font_cn(run)


def add_blockquote_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = p.add_run(part[2:-2])
            run.bold = True
            run.italic = True
            _set_run_font_cn(run)
        else:
            run = p.add_run(part)
            run.italic = True
            _set_run_font_cn(run)


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


def add_table_from_md(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.rows[i].cells[j]
            plain = row[j] if j < len(row) else ""
            plain = re.sub(r"\*\*([^*]+)\*\*", r"\1", plain)
            cell.text = plain
            for p in cell.paragraphs:
                for r in p.runs:
                    _set_run_font_cn(r, size_pt=9)


def md_to_docx(md_text: str, out_path: Path) -> None:
    doc = Document()
    p_note = doc.add_paragraph()
    r_note = p_note.add_run(
        "说明：本文件由 Markdown 转换生成。源文件：google play/docs/business_design/business_scope_and_metrics_v2.md"
    )
    _set_run_font_cn(r_note, size_pt=9)
    r_note.italic = True

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []
    table_buf: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_buf
        if not table_buf:
            return
        cleaned = []
        for row in table_buf:
            joined = "|".join(row)
            core = joined.replace("|", "").replace("-", "").replace(":", "").strip()
            if not core:
                continue
            cleaned.append(row)
        if cleaned:
            add_table_from_md(doc, cleaned)
        table_buf = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                flush_table()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()
        if stripped == "---":
            flush_table()
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            table_buf.append(cells)
            i += 1
            continue
        else:
            flush_table()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("> "):
            inner = stripped[2:].strip()
            add_blockquote_paragraph(doc, inner)
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            title = re.sub(r"\*\*([^*]+)\*\*", r"\1", m.group(2).strip())
            doc.add_heading(title, level=min(level, 9))
            i += 1
            continue

        if stripped.startswith(("- ", "* ")):
            text = stripped[2:].strip()
            try:
                add_paragraph_with_bold(doc, text, style="List Bullet")
            except KeyError:
                add_paragraph_with_bold(doc, "• " + text)
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            try:
                add_paragraph_with_bold(doc, re.sub(r"^\d+\.\s", "", stripped), style="List Number")
            except KeyError:
                add_paragraph_with_bold(doc, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue

        add_paragraph_with_bold(doc, stripped)
        i += 1

    flush_table()
    doc.save(out_path)


def main() -> int:
    if not MD_PATH.is_file():
        print(f"Missing {MD_PATH}", file=sys.stderr)
        return 1
    text = MD_PATH.read_text(encoding="utf-8")
    md_to_docx(text, DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
